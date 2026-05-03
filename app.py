import os
import re
import uuid
from collections import Counter
from pathlib import Path

from flask import Flask, render_template, request
from pdfminer.high_level import extract_text as extract_pdf_text
from werkzeug.utils import secure_filename

try:
    from docx import Document
except ImportError:  # DOCX support is optional; PDF and TXT still work.
    Document = None


BASE_DIR = Path(__file__).resolve().parent
UPLOAD_DIR = BASE_DIR / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

ALLOWED_EXTENSIONS = {"pdf", "txt", "docx"}

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 6 * 1024 * 1024


SKILL_CATEGORIES = {
    "Programming": [
        "python", "java", "javascript", "typescript", "c", "c++", "c#", "go",
        "php", "ruby", "kotlin", "swift", "sql", "html", "css",
    ],
    "Frameworks": [
        "flask", "django", "fastapi", "spring", "spring boot", "react",
        "angular", "vue", "node.js", "express", "bootstrap", "tailwind",
    ],
    "Data & AI": [
        "machine learning", "deep learning", "artificial intelligence", "nlp",
        "data analysis", "data science", "pandas", "numpy", "scikit-learn",
        "tensorflow", "pytorch", "power bi", "tableau", "excel",
    ],
    "Databases": [
        "mysql", "postgresql", "mongodb", "sqlite", "oracle", "redis",
        "firebase", "nosql",
    ],
    "Cloud & DevOps": [
        "aws", "azure", "gcp", "docker", "kubernetes", "git", "github",
        "ci/cd", "linux", "jenkins",
    ],
    "Core CS": [
        "data structures", "algorithms", "oop", "object oriented programming",
        "api", "rest api", "microservices", "system design",
    ],
    "Professional": [
        "communication", "leadership", "teamwork", "problem solving",
        "collaboration", "project management", "analytical thinking",
    ],
}

ALL_SKILLS = sorted({skill for skills in SKILL_CATEGORIES.values() for skill in skills})

ACTION_VERBS = {
    "achieved", "automated", "built", "created", "delivered", "designed",
    "developed", "improved", "increased", "launched", "led", "managed",
    "optimized", "reduced", "resolved", "scaled", "streamlined",
}

SECTION_ALIASES = {
    "Summary": ["summary", "profile", "objective", "career objective"],
    "Skills": ["skills", "technical skills", "core competencies"],
    "Experience": ["experience", "work experience", "employment", "internship"],
    "Education": ["education", "academic background", "qualification"],
    "Projects": ["projects", "academic projects", "personal projects"],
    "Certifications": ["certifications", "certificates", "licenses"],
}

STOPWORDS = {
    "and", "the", "for", "with", "that", "this", "from", "your", "you",
    "are", "will", "our", "has", "have", "job", "role", "work", "team",
    "using", "use", "into", "their", "they", "candidate", "experience",
    "skills", "ability", "knowledge", "responsibilities", "requirements",
}


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def normalize_text(text):
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_docx_text(file_path):
    if Document is None:
        raise RuntimeError("DOCX support requires python-docx. Upload a PDF/TXT or install python-docx.")

    document = Document(file_path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_text = []
    for table in document.tables:
        for row in table.rows:
            table_text.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(paragraphs + table_text)


def extract_resume_text(file_path):
    extension = Path(file_path).suffix.lower()
    if extension == ".pdf":
        return normalize_text(extract_pdf_text(file_path))
    if extension == ".txt":
        return normalize_text(Path(file_path).read_text(encoding="utf-8", errors="ignore"))
    if extension == ".docx":
        return normalize_text(extract_docx_text(file_path))
    raise ValueError("Unsupported file type.")


def find_contact_details(text):
    email = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)
    phone = re.search(r"(\+?\d[\d\s().-]{8,}\d)", text)
    linkedin = re.search(r"(linkedin\.com/in/[A-Za-z0-9_-]+)", text, re.IGNORECASE)
    portfolio = re.search(r"(github\.com/[A-Za-z0-9_-]+|https?://[^\s]+)", text, re.IGNORECASE)

    return {
        "email": email.group(0) if email else None,
        "phone": phone.group(0).strip() if phone else None,
        "linkedin": linkedin.group(0) if linkedin else None,
        "portfolio": portfolio.group(0) if portfolio else None,
    }


def find_sections(text):
    lower_text = text.lower()
    found = []
    missing = []

    for section, aliases in SECTION_ALIASES.items():
        section_found = any(re.search(rf"(^|\n)\s*{re.escape(alias)}\s*[:\n-]", lower_text) for alias in aliases)
        if not section_found:
            section_found = any(alias in lower_text for alias in aliases)

        if section_found:
            found.append(section)
        else:
            missing.append(section)

    return found, missing


def extract_skills(text):
    lower_text = text.lower()
    detected = {}

    for category, skills in SKILL_CATEGORIES.items():
        matches = []
        for skill in skills:
            pattern = rf"(?<![a-z0-9+#]){re.escape(skill)}(?![a-z0-9+#])"
            if re.search(pattern, lower_text):
                matches.append(skill)
        if matches:
            detected[category] = sorted(set(matches))

    flat = sorted({skill for skills in detected.values() for skill in skills})
    return detected, flat


def extract_job_keywords(job_description):
    if not job_description:
        return []

    lower_description = job_description.lower()
    skill_hits = [skill for skill in ALL_SKILLS if skill in lower_description]

    words = re.findall(r"[a-z][a-z+#.-]{2,}", lower_description)
    useful_words = [
        word for word in words
        if word not in STOPWORDS and not word.isdigit() and len(word) >= 3
    ]
    frequent_words = [word for word, _ in Counter(useful_words).most_common(20)]

    return sorted(set(skill_hits + frequent_words))


def match_keywords(text, job_description):
    keywords = extract_job_keywords(job_description)
    if not keywords:
        return [], []

    lower_text = text.lower()
    matched = [keyword for keyword in keywords if keyword in lower_text]
    missing = [keyword for keyword in keywords if keyword not in lower_text]
    return matched, missing


def count_action_verbs(text):
    words = re.findall(r"[a-zA-Z]+", text.lower())
    return sum(1 for word in words if word in ACTION_VERBS)


def count_quantified_results(text):
    metric_patterns = [
        r"\b\d+%",
        r"\b\d+\+",
        r"\b\d+x\b",
        r"\b\d+\s*(users|clients|customers|students|projects|apis|reports|dashboards|hours|days|weeks|months)\b",
        r"[$₹€]\s?\d+",
    ]
    return sum(len(re.findall(pattern, text, re.IGNORECASE)) for pattern in metric_patterns)


def score_resume(text, job_description=""):
    words = re.findall(r"\b[\w+#.-]+\b", text)
    word_count = len(words)
    contact = find_contact_details(text)
    found_sections, missing_sections = find_sections(text)
    categorized_skills, flat_skills = extract_skills(text)
    matched_keywords, missing_keywords = match_keywords(text, job_description)
    action_verb_count = count_action_verbs(text)
    metric_count = count_quantified_results(text)
    line_count = max(1, len([line for line in text.splitlines() if line.strip()]))
    bullet_count = len(re.findall(r"(^|\n)\s*([•*-]|\d+[.)])\s+", text))

    contact_score = 0
    contact_score += 5 if contact["email"] else 0
    contact_score += 5 if contact["phone"] else 0
    contact_score += 3 if contact["linkedin"] else 0
    contact_score += 2 if contact["portfolio"] else 0

    section_score = round((len(found_sections) / len(SECTION_ALIASES)) * 20)

    if job_description:
        keyword_score = round((len(matched_keywords) / max(1, len(matched_keywords) + len(missing_keywords))) * 25)
    else:
        keyword_score = min(25, round((len(flat_skills) / 14) * 25))

    impact_score = min(20, metric_count * 3 + action_verb_count)

    readability_score = 0
    readability_score += 4 if bullet_count >= 4 else max(0, bullet_count)
    readability_score += 3 if 350 <= word_count <= 950 else 1
    readability_score += 3 if word_count / line_count <= 30 else 1

    length_score = 10 if 350 <= word_count <= 950 else 7 if 220 <= word_count < 350 or 950 < word_count <= 1200 else 4

    total_score = min(100, contact_score + section_score + keyword_score + impact_score + readability_score + length_score)

    breakdown = [
        {
            "name": "Contact & Links",
            "score": contact_score,
            "max": 15,
            "feedback": "Recruiter contact details are easy to find." if contact_score >= 10 else "Add email, phone, LinkedIn, and portfolio/GitHub links near the top.",
        },
        {
            "name": "Resume Sections",
            "score": section_score,
            "max": 20,
            "feedback": f"Found: {', '.join(found_sections) or 'none'}",
        },
        {
            "name": "Keyword Match",
            "score": keyword_score,
            "max": 25,
            "feedback": "Compared with the job description." if job_description else "Scored from detected industry skills because no job description was provided.",
        },
        {
            "name": "Impact & Achievements",
            "score": impact_score,
            "max": 20,
            "feedback": f"Found {metric_count} quantified result(s) and {action_verb_count} strong action verb(s).",
        },
        {
            "name": "ATS Readability",
            "score": readability_score,
            "max": 10,
            "feedback": "Bullet structure and line length look ATS-friendly." if readability_score >= 7 else "Use concise bullets and avoid dense paragraphs.",
        },
        {
            "name": "Length",
            "score": length_score,
            "max": 10,
            "feedback": f"Detected approximately {word_count} words.",
        },
    ]

    suggestions = build_suggestions(
        contact=contact,
        missing_sections=missing_sections,
        flat_skills=flat_skills,
        missing_keywords=missing_keywords,
        metric_count=metric_count,
        action_verb_count=action_verb_count,
        bullet_count=bullet_count,
        word_count=word_count,
        has_job_description=bool(job_description),
    )

    return {
        "score": total_score,
        "grade": grade_for_score(total_score),
        "word_count": word_count,
        "contact": contact,
        "sections_found": found_sections,
        "sections_missing": missing_sections,
        "skills_by_category": categorized_skills,
        "skills": flat_skills,
        "matched_keywords": matched_keywords[:30],
        "missing_keywords": missing_keywords[:30],
        "breakdown": breakdown,
        "suggestions": suggestions,
        "preview": text[:1200],
    }


def grade_for_score(score):
    if score >= 85:
        return "Excellent"
    if score >= 70:
        return "Strong"
    if score >= 55:
        return "Needs polish"
    return "Needs work"


def build_suggestions(
    contact,
    missing_sections,
    flat_skills,
    missing_keywords,
    metric_count,
    action_verb_count,
    bullet_count,
    word_count,
    has_job_description,
):
    suggestions = []

    missing_contact = [
        label for key, label in {
            "email": "email",
            "phone": "phone number",
            "linkedin": "LinkedIn URL",
            "portfolio": "GitHub or portfolio link",
        }.items()
        if not contact[key]
    ]
    if missing_contact:
        suggestions.append({
            "title": "Make recruiter contact effortless",
            "detail": f"Add your {', '.join(missing_contact)} in the header so ATS systems and recruiters can parse it.",
        })

    important_missing = [section for section in missing_sections if section in {"Summary", "Skills", "Experience", "Education", "Projects"}]
    if important_missing:
        suggestions.append({
            "title": "Add missing standard sections",
            "detail": f"Include clear headings for {', '.join(important_missing)}. Simple headings improve ATS parsing.",
        })

    if len(flat_skills) < 8:
        suggestions.append({
            "title": "Strengthen the skills section",
            "detail": "List 8-15 relevant tools, languages, frameworks, databases, and domain skills using exact terms from target jobs.",
        })

    if has_job_description and missing_keywords:
        suggestions.append({
            "title": "Tailor keywords to the job",
            "detail": f"Consider adding truthful examples for these missing terms: {', '.join(missing_keywords[:8])}.",
        })

    if metric_count < 3:
        suggestions.append({
            "title": "Quantify your achievements",
            "detail": "Add numbers such as percentage improvements, users served, reports built, time saved, revenue, or project scale.",
        })

    if action_verb_count < 6:
        suggestions.append({
            "title": "Start bullets with stronger verbs",
            "detail": "Use verbs like built, automated, optimized, led, improved, reduced, delivered, and launched.",
        })

    if bullet_count < 4:
        suggestions.append({
            "title": "Use concise bullet points",
            "detail": "ATS tools handle clean bullets better than long paragraphs. Keep each bullet focused on one result.",
        })

    if word_count < 250:
        suggestions.append({
            "title": "Add more role-specific evidence",
            "detail": "The resume looks short. Add projects, internships, responsibilities, achievements, and technology details.",
        })
    elif word_count > 1100:
        suggestions.append({
            "title": "Trim for recruiter scanning",
            "detail": "The resume is long. Keep recent and relevant details, remove repetition, and tighten older experience.",
        })

    if not suggestions:
        suggestions.append({
            "title": "Resume is in strong shape",
            "detail": "Keep tailoring the summary, skills, and top bullets to each specific job description.",
        })

    return suggestions


@app.route("/")
def home():
    return render_template("index.html")


@app.route("/analyze", methods=["POST"])
def analyze():
    resume_file = request.files.get("resume")
    job_description = request.form.get("job_description", "").strip()

    if not resume_file or not resume_file.filename:
        return render_template("index.html", error="Please upload a resume file.")

    if not allowed_file(resume_file.filename):
        return render_template("index.html", error="Upload a PDF, DOCX, or TXT resume.")

    original_name = secure_filename(resume_file.filename)
    extension = Path(original_name).suffix.lower()
    saved_path = UPLOAD_DIR / f"{uuid.uuid4().hex}{extension}"

    try:
        resume_file.save(saved_path)
        text = extract_resume_text(saved_path)
        if len(text) < 80:
            return render_template(
                "index.html",
                error="I could not read enough text from that file. Try a text-based PDF, DOCX, or TXT resume.",
            )

        analysis = score_resume(text, job_description)
        return render_template("result.html", analysis=analysis, filename=original_name)
    except Exception as exc:
        return render_template("index.html", error=f"Could not analyze the resume: {exc}")
    finally:
        if saved_path.exists():
            saved_path.unlink()


if __name__ == "__main__":
    debug = os.environ.get("FLASK_DEBUG", "").lower() in {"1", "true", "yes"}
    app.run(debug=debug)
