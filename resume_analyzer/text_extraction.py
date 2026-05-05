"""File validation and resume text extraction helpers."""

import re
from pathlib import Path

ALLOWED_EXTENSIONS = {"pdf", "txt", "docx"}


def allowed_file(filename):
    """Return True when the upload filename has a supported extension."""
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def normalize_text(text):
    """Normalize extracted resume text without changing its semantic content."""
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_docx_text(file_path):
    """Extract paragraphs and table cells from a DOCX file."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "DOCX support requires python-docx. Upload a PDF/TXT or install python-docx."
        ) from exc

    document = Document(file_path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_text = []
    for table in document.tables:
        for row in table.rows:
            table_text.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(paragraphs + table_text)


def extract_resume_text(file_path):
    """Extract normalized text from a PDF, TXT, or DOCX resume."""
    extension = Path(file_path).suffix.lower()
    if extension == ".pdf":
        from pdfminer.high_level import extract_text as extract_pdf_text

        return normalize_text(extract_pdf_text(file_path))
    if extension == ".txt":
        return normalize_text(Path(file_path).read_text(encoding="utf-8", errors="ignore"))
    if extension == ".docx":
        return normalize_text(extract_docx_text(file_path))
    raise ValueError("Unsupported file type.")
